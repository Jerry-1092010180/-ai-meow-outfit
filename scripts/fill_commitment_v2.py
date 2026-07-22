from docx import Document
from docx.shared import Pt

doc = Document('/Users/jerry/Downloads/OPC大赛-原创与AI使用承诺书.docx')

# ── 基本信息 ──
paragraphs = doc.paragraphs

if '参赛者/团队名称' in paragraphs[2].text:
    paragraphs[2].clear()
    paragraphs[2].add_run('参赛者/团队名称：Jerry（个人参赛）')

if '参赛项目名称' in paragraphs[3].text:
    paragraphs[3].clear()
    paragraphs[3].add_run('参赛项目名称：AI喵搭')

if '参赛编号' in paragraphs[4].text:
    paragraphs[4].clear()
    paragraphs[4].add_run('参赛编号：待组委会分配')

if '参赛赛道' in paragraphs[5].text:
    paragraphs[5].clear()
    paragraphs[5].add_run('参赛赛道：□ 自由赛道　☑ 企业赛道（Bounty）')

if '子方向' in paragraphs[6].text:
    paragraphs[6].clear()
    paragraphs[6].add_run('子方向：☑ 创想家　□ 智造家')

if 'Bounty赛题' in paragraphs[7].text:
    paragraphs[7].clear()
    paragraphs[7].add_run('Bounty赛题（如适用）：03 - 为4500万会员设计一款AI原生的喵街互动玩法（银泰商业）')

# ── AI 工具表格 ──
table = doc.tables[0]
# Add data row
row = table.add_row()
row.cells[0].paragraphs[0].clear()
row.cells[0].paragraphs[0].add_run('Claude Code (Claude Opus 4.8)')
row.cells[1].paragraphs[0].clear()
row.cells[1].paragraphs[0].add_run('全栈开发：产品设计、前端（React/Three.js）开发、后端（Python/FastAPI）开发、3D渲染引擎、文档撰写')
row.cells[2].paragraphs[0].clear()
row.cells[2].paragraphs[0].add_run('代码生成、架构建议、技术方案设计、Bug修复、文档生成')
row.cells[3].paragraphs[0].clear()
row.cells[3].paragraphs[0].add_run('约60%（核心创意与商业判断为本人独立完成）')

for cell in row.cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)

# ── 签名日期（第一个日期单元格） ──
sign_table = doc.tables[1]
date_cell = sign_table.rows[1].cells[2]
date_cell.paragraphs[0].clear()
date_cell.paragraphs[0].add_run('2026年7月18日')

out = '/Users/jerry/Downloads/OPC大赛-原创与AI使用承诺书_已填写.docx'
doc.save(out)
print(f'Saved to {out}')
