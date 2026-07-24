from docx import Document
from docx.shared import Pt, RGBColor
import copy

doc = Document('/Users/jerry/Downloads/OPC大赛-原创与AI使用承诺书.docx')

# Find and fill form fields
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                text = p.text.strip()

                # Fill project info
                if '参赛者/团队名称' in text and '____' in text:
                    p.clear()
                    p.add_run('参赛者/团队名称：Jerry（个人参赛）')

                elif '参赛项目名称' in text and '____' in text:
                    p.clear()
                    p.add_run('参赛项目名称：AI喵搭')

                elif '参赛编号' in text and '____' in text:
                    p.clear()
                    p.add_run('参赛编号：待组委会分配')

                elif '参赛赛道' in text or '企业赛道' in text or 'Bounty' in text:
                    if '□' in text or '自由赛道' in text:
                        p.clear()
                        run = p.add_run('参赛赛道：□ 自由赛道　☑ 企业赛道（Bounty）')

                elif '子方向' in text:
                    if '□' in text or '创想家' in text:
                        p.clear()
                        run = p.add_run('子方向：☑ 创想家　□ 智造家')

                elif 'Bounty赛题' in text and '____' in text:
                    p.clear()
                    run = p.add_run('Bounty赛题（如适用）：03 - 为4500万会员设计一款AI原生的喵街互动玩法（银泰商业）')

# Fill AI tool disclosure table
for table in doc.tables:
    rows = table.rows
    if len(rows) >= 3:
        # Check if this is the AI tools table (has headers: AI工具/模型名称, 使用环节, etc.)
        header_text = rows[0].cells[0].text if len(rows[0].cells) > 0 else ''
        if 'AI工具' in header_text or '模型' in header_text:
            # Add data row if not already filled
            if len(rows) == 2:  # Only header row exists
                from docx.oxml.ns import qn
                # Add a row
                row = table.add_row()
                row.cells[0].paragraphs[0].clear()
                row.cells[0].paragraphs[0].add_run('Claude Code (Claude Opus 4.8)')
                row.cells[1].paragraphs[0].clear()
                row.cells[1].paragraphs[0].add_run('全栈开发：产品设计、前端（React/Three.js）开发、后端（Python/FastAPI）开发、3D渲染引擎、文档撰写')
                row.cells[2].paragraphs[0].clear()
                row.cells[2].paragraphs[0].add_run('代码生成、架构建议、技术方案设计、Bug修复、文档生成')
                row.cells[3].paragraphs[0].clear()
                row.cells[3].paragraphs[0].add_run('约60%（核心创意与商业判断为本人独立完成）')

# Set font to 10pt for readability on all filled cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if run.font.size is None or run.font.size < Pt(9):
                        run.font.size = Pt(10)
                    run.font.name = '微软雅黑'

# Fill signature section - date
# Search for date placeholders like "年    月    日"
for p in doc.paragraphs:
    text = p.text
    if '年' in text and '月' in text and '日' in text and len(text) < 20 and ('____' in text or '  ' in text):
        # Check if this looks like a date line (not a template description)
        if '年' in text and '月' in text and '日' in text:
            p.clear()
            p.add_run('2026 年  7 月  18 日')

doc.save('/Users/jerry/Downloads/OPC大赛-原创与AI使用承诺书_已填写.docx')
print("DOCX saved successfully")
